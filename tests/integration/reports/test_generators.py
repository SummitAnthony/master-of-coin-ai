"""Integration tests: the three binary generators + the pack orchestrator."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from advisor.narrative.advisor import Advisory
from advisor.reports.excel import write_excel
from advisor.reports.pack import generate_report_pack
from advisor.reports.pdf import build_pdf_bytes, write_pdf
from advisor.reports.word import write_word
from advisor.schema import Facts


def test_excel_roundtrip(
    facts: Facts, advisory: Advisory, gen_at: datetime, tmp_path: Path
) -> None:
    path = write_excel(facts, advisory, tmp_path / "r.xlsx", generated_at=gen_at)
    wb = load_workbook(path)
    assert {"Cover", "Income Statement", "KPIs", "Status", "Variance", "Anomalies"} <= set(
        wb.sheetnames
    )
    income = wb["Income Statement"]
    assert income.cell(row=1, column=2).value == "FY2022-23"
    assert income.cell(row=2, column=2).value == float(facts.kpis[0].revenue)
    wb.close()
    path.unlink()  # Windows file-lock regression guard


def test_word_roundtrip(facts: Facts, advisory: Advisory, gen_at: datetime, tmp_path: Path) -> None:
    path = write_word(facts, advisory, tmp_path / "r.docx", generated_at=gen_at)
    doc = Document(str(path))
    headings = [p.text for p in doc.paragraphs]
    assert "Executive Summary" in headings
    assert "Recommendations" in headings
    assert len(doc.tables) >= 1


def test_pdf_bytes_are_valid(
    facts: Facts, advisory: Advisory, gen_at: datetime, tmp_path: Path
) -> None:
    data = build_pdf_bytes(facts, advisory, generated_at=gen_at)
    assert data.startswith(b"%PDF-")
    assert b"%%EOF" in data
    path = write_pdf(facts, advisory, tmp_path / "r.pdf", generated_at=gen_at)
    assert path.stat().st_size > 1000


def test_pdf_is_deterministic(facts: Facts, advisory: Advisory, gen_at: datetime) -> None:
    a = build_pdf_bytes(facts, advisory, generated_at=gen_at)
    b = build_pdf_bytes(facts, advisory, generated_at=gen_at)
    assert a == b


def test_pack_generates_four_files(
    facts: Facts, advisory: Advisory, gen_at: datetime, tmp_path: Path
) -> None:
    pack = generate_report_pack(facts, advisory, tmp_path / "out", generated_at=gen_at)
    for p in (pack.excel_path, pack.pdf_path, pack.word_path, pack.dashboard_path):
        assert p.exists() and p.stat().st_size > 0
