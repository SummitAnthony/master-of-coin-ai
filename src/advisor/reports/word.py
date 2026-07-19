"""Branded Word memo (python-docx)."""

from __future__ import annotations

from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document as _new_document

from advisor import PRODUCT_NAME
from advisor.narrative.advisor import Advisory
from advisor.schema import EntityMeta, Facts

from .formatting import DISCLAIMER, format_money, format_pct, format_per_unit, format_volume

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

# (row label — "{unit}" is replaced with the entity's volume unit,
#  PeriodKPIs attribute, format kind)
_KPI_ROWS: tuple[tuple[str, str, str], ...] = (
    ("Revenue", "revenue", "money"),
    ("Gross profit", "gross_profit", "money"),
    ("Net profit", "net_profit", "money"),
    ("Gross margin", "gross_margin_pct", "pct"),
    ("Operating margin", "operating_margin_pct", "pct"),
    ("Net margin", "net_margin_pct", "pct"),
    ("Volume", "volume_mt", "volume"),
    ("Selling price / {unit}", "selling_price_per_mt", "per_unit"),
)


def _fmt(value: Decimal | None, kind: str, entity: EntityMeta) -> str:
    if kind == "money":
        return format_money(value, currency=entity.currency)
    if kind == "pct":
        return format_pct(value)
    if kind == "volume":
        return format_volume(value, unit=entity.volume_unit)
    return format_per_unit(value, currency=entity.currency, unit=entity.volume_unit)


def build_document(facts: Facts, narrative: Advisory, *, generated_at: datetime) -> DocxDocument:
    doc = _new_document()
    entity = facts.entity
    title = (
        f"{entity.group_name} — {entity.company_name}" if entity.group_name else entity.company_name
    )
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"{PRODUCT_NAME} — Financial Advisory Memo")
    doc.add_paragraph(f"Generated: {generated_at.isoformat()}")
    doc.add_paragraph(f"Latest period: {facts.latest_period.label}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(narrative.executive_summary)

    doc.add_heading("Key Performance Indicators", level=1)
    table = doc.add_table(rows=1, cols=1 + len(facts.kpis))
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Metric"
    for i, k in enumerate(facts.kpis):
        header[1 + i].text = k.period.label
    for label, attr, kind in _KPI_ROWS:
        cells = table.add_row().cells
        cells[0].text = label.replace("{unit}", entity.volume_unit)
        for i, k in enumerate(facts.kpis):
            cells[1 + i].text = _fmt(getattr(k, attr), kind, entity)

    doc.add_heading("Risk Commentary", level=1)
    doc.add_paragraph(narrative.risk_commentary)

    doc.add_heading("Recommendations", level=1)
    for rec in narrative.recommendations:
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_heading("Anomalies", level=1)
    if facts.anomalies:
        for a in facts.anomalies:
            doc.add_paragraph(f"[{a.severity.value}] {a.code} — {a.period} ({a.metric})")
    else:
        doc.add_paragraph("No anomalies detected.")

    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(DISCLAIMER)
    return doc


def write_word(facts: Facts, narrative: Advisory, path: Path, *, generated_at: datetime) -> Path:
    build_document(facts, narrative, generated_at=generated_at).save(str(path))
    return path
